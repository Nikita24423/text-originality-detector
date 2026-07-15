"""
Извлечение текста из документов: абзацы, таблицы, OCR с изображений.

Поддерживаемые форматы: .docx, .pdf
"""
from __future__ import annotations

import os
import zipfile
from dataclasses import dataclass
from io import BytesIO
from typing import List, Optional

from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


@dataclass
class ExtractionStats:
    paragraphs: int = 0
    tables: int = 0
    images_ocr: int = 0
    pages: int = 0
    ocr_available: bool = False

    def summary(self) -> str:
        parts: List[str] = []
        if self.pages:
            parts.append(f"{self.pages} стр.")
        if self.paragraphs:
            parts.append(f"{self.paragraphs} абз.")
        if self.tables:
            parts.append(f"{self.tables} табл.")
        if self.images_ocr:
            parts.append(f"{self.images_ocr} OCR")
        return ", ".join(parts) if parts else "текст"


_ocr_engine = None
_ocr_checked = False


def _get_ocr_engine():
    global _ocr_engine, _ocr_checked
    if _ocr_checked:
        return _ocr_engine
    _ocr_checked = True
    try:
        from rapidocr_onnxruntime import RapidOCR

        _ocr_engine = RapidOCR()
    except Exception:
        _ocr_engine = None
    return _ocr_engine


def ocr_available() -> bool:
    return _get_ocr_engine() is not None


def _ocr_image_bytes(image_data: bytes, min_size: int = 80) -> str:
    engine = _get_ocr_engine()
    if engine is None:
        return ""

    try:
        from PIL import Image
        import numpy as np

        image = Image.open(BytesIO(image_data)).convert("RGB")
        if image.width < min_size or image.height < min_size:
            return ""

        result, _ = engine(np.array(image))
        if not result:
            return ""

        lines = [str(item[1]).strip() for item in result if len(item) > 1 and str(item[1]).strip()]
        return "\n".join(lines)
    except Exception:
        return ""


def _table_to_text(table: Table) -> str:
    rows: List[str] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        cells = [c for c in cells if c]
        if cells:
            rows.append(" | ".join(cells))
    return "\n".join(rows)


def _iter_docx_blocks(document):
    parent_elm = document.element.body
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _extract_docx_images(content: bytes, stats: ExtractionStats) -> List[str]:
    parts: List[str] = []
    ocr_ext = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tif", ".tiff", ".webp"}

    with zipfile.ZipFile(BytesIO(content)) as archive:
        for name in archive.namelist():
            if not name.startswith("word/media/"):
                continue
            ext = os.path.splitext(name.lower())[1]
            if ext not in ocr_ext:
                continue

            text = _ocr_image_bytes(archive.read(name))
            if not text:
                continue

            stats.images_ocr += 1
            image_name = os.path.basename(name)
            parts.append(f"[Текст с изображения: {image_name}]\n{text}")

    return parts


def extract_docx(content: bytes) -> tuple[str, ExtractionStats]:
    from docx import Document

    stats = ExtractionStats(ocr_available=ocr_available())
    document = Document(BytesIO(content))
    parts: List[str] = []
    table_index = 0

    for block in _iter_docx_blocks(document):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                stats.paragraphs += 1
                parts.append(text)
        elif isinstance(block, Table):
            table_text = _table_to_text(block)
            if table_text:
                table_index += 1
                stats.tables += 1
                parts.append(f"[Таблица {table_index}]\n{table_text}")

    image_parts = _extract_docx_images(content, stats)
    parts.extend(image_parts)

    return "\n\n".join(parts).strip(), stats


def _pdf_tables_to_text(tables: list) -> List[str]:
    parts: List[str] = []
    for table_index, table in enumerate(tables, 1):
        rows: List[str] = []
        for row in table:
            cells = [str(cell).strip().replace("\n", " ") for cell in row if cell]
            cells = [c for c in cells if c]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            parts.append(f"[Таблица {table_index}]\n" + "\n".join(rows))
    return parts


def extract_pdf(content: bytes) -> tuple[str, ExtractionStats]:
    import pdfplumber

    stats = ExtractionStats(ocr_available=ocr_available())
    parts: List[str] = []

    with pdfplumber.open(BytesIO(content)) as pdf:
        stats.pages = len(pdf.pages)
        for page_number, page in enumerate(pdf.pages, 1):
            page_parts: List[str] = []

            text = (page.extract_text() or "").strip()
            if text:
                stats.paragraphs += 1
                page_parts.append(text)

            tables = page.extract_tables() or []
            table_texts = _pdf_tables_to_text(tables)
            if table_texts:
                stats.tables += len(table_texts)
                page_parts.extend(table_texts)

            if not page_parts:
                try:
                    page_image = page.to_image(resolution=200)
                    buffer = BytesIO()
                    page_image.original.save(buffer, format="PNG")
                    ocr_text = _ocr_image_bytes(buffer.getvalue(), min_size=120)
                    if ocr_text:
                        stats.images_ocr += 1
                        page_parts.append(f"[Текст со страницы {page_number} (OCR)]\n{ocr_text}")
                except Exception:
                    pass

            if page_parts:
                parts.append(f"--- Страница {page_number} ---\n" + "\n\n".join(page_parts))

    return "\n\n".join(parts).strip(), stats
